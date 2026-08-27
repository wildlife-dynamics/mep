"""Shared fixtures for ecoscope_workflows_ext_mep.tasks.reporting tests.

The function under test (`generate_subject_report`) is registered via
`wt_registry.register()`, a no-op decorator at call time, so it is called
directly as plain Python -- no workflow engine involved.

Fixtures here focus on building minimal, real inputs (docx templates via
python-docx, images via Pillow) rather than mocking the document-generation
libraries themselves, since docxtpl / Pillow are all installed in this
environment and cheap to run for real.
"""

from __future__ import annotations

from pathlib import Path
from typing import Callable

import docx
import pytest
from PIL import Image


@pytest.fixture
def make_png(tmp_path) -> Callable[..., Path]:
    """Factory that writes a small, valid PNG at a given path under tmp_path."""

    def _make(path: Path, size: tuple[int, int] = (32, 32), color=(255, 0, 0)) -> Path:
        path.parent.mkdir(parents=True, exist_ok=True)
        Image.new("RGB", size, color=color).save(path)
        return path

    return _make


@pytest.fixture
def make_docx_template(tmp_path) -> Callable[..., Path]:
    """Factory that writes a minimal .docx with the given paragraph strings
    (each may contain Jinja/docxtpl placeholders like '{{ field }}') and
    returns its path.
    """

    counter = {"n": 0}

    def _make(paragraphs: list[str], name: str | None = None) -> Path:
        counter["n"] += 1
        filename = name or f"template_{counter['n']}.docx"
        path = tmp_path / filename
        path.parent.mkdir(parents=True, exist_ok=True)
        doc = docx.Document()
        for text in paragraphs:
            doc.add_paragraph(text)
        doc.save(path)
        return path

    return _make


@pytest.fixture
def read_docx_text() -> Callable[[Path], list[str]]:
    """Return the paragraph texts of a saved .docx, for asserting render output."""

    def _read(path: Path) -> list[str]:
        d = docx.Document(str(path))
        return [p.text for p in d.paragraphs]

    return _read
